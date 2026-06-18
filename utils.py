import openpyxl
import os
import json
import requests
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import parse_xml, OxmlElement
from docx.oxml.ns import nsdecls, qn

def format_id(num):
    if num is None:
        return "-"
    try:
        return f"{int(num):,}".replace(",", ".")
    except:
        return str(num)

def set_cell_shading(cell, color_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = tcPr.find(qn('w:shd'))
    if shd is not None:
        tcPr.remove(shd)
    shading_xml = f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>'
    tcPr.append(parse_xml(shading_xml))

def set_table_borders(table, color="D3D3D3", sz="4", val="single"):
    tblPr = table._tbl.tblPr
    borders_xml = f'''
    <w:tblBorders {nsdecls("w")}>
        <w:top w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>
        <w:bottom w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>
        <w:left w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>
        <w:right w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>
        <w:insideH w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>
        <w:insideV w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>
    </w:tblBorders>
    '''
    tblBorders = tblPr.find(qn('w:tblBorders'))
    if tblBorders is not None:
        tblPr.remove(tblBorders)
    tblPr.append(parse_xml(borders_xml))

def set_table_margins(table, top=80, bottom=80, left=120, right=120):
    tblPr = table._tbl.tblPr
    tblCellMar = tblPr.find(qn('w:tblCellMar'))
    if tblCellMar is not None:
        tblPr.remove(tblCellMar)
    tblCellMar = OxmlElement('w:tblCellMar')
    for margin, val in [('w:top', top), ('w:bottom', bottom), ('w:left', left), ('w:right', right)]:
        node = OxmlElement(margin)
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tblCellMar.append(node)
    tblPr.append(tblCellMar)

def format_run(run, font_name="Arial", size_pt=10, color_rgb=None, bold=False, italic=False):
    run.font.name = font_name
    run.font.size = Pt(size_pt)
    run.bold = bold
    run.italic = italic
    if color_rgb:
        run.font.color.rgb = color_rgb
    
    rPr = run._r.get_or_add_rPr()
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:ascii'), font_name)
    rFonts.set(qn('w:hAnsi'), font_name)
    rFonts.set(qn('w:cs'), font_name)
    rPr.append(rFonts)

def set_para_spacing(paragraph, before_pt=0, after_pt=4, line_spacing=1.15):
    paragraph.paragraph_format.space_before = Pt(before_pt)
    paragraph.paragraph_format.space_after = Pt(after_pt)
    paragraph.paragraph_format.line_spacing = line_spacing

def load_token_limits_db(pricing_xlsx_path):
    if not pricing_xlsx_path or not os.path.exists(pricing_xlsx_path):
        # Fallback dictionary of default limits
        return {}, {}
    
    try:
        wb = openpyxl.load_workbook(pricing_xlsx_path, data_only=True)
        sheet = wb.active
        
        dept_cols = {}
        for col in range(3, 10):
            val = sheet.cell(row=5, column=col).value
            if val:
                dept_cols[val.lower().strip()] = col
                
        dept_mapping = {
            "collection": "collection",
            "credit data analyst": "credit",
            "credit agent": "credit",
            "corporate university": "corpu",
            "ssd": "ssd",
            "telecenter": "telecenter",
            "marketing": "marketing"
        }
        
        limits = {}
        current_section = None
        for r in range(4, 50):
            val_a = sheet.cell(row=r, column=1).value
            if val_a and isinstance(val_a, str):
                if "15 MENIT" in val_a.upper():
                    current_section = "15m"
                elif "12 JAM" in val_a.upper():
                    current_section = "12j"
                elif "24 JAM" in val_a.upper():
                    current_section = "24j"
                
            model_cell = sheet.cell(row=r, column=1).value
            if model_cell and isinstance(model_cell, str) and model_cell.strip() in ['Sonic', 'Cheetah', 'Meticulous']:
                model_base = model_cell.strip()
                row_input = r + 1
                row_output = r + 2
                
                for dept_name, col_idx in dept_cols.items():
                    in_val = sheet.cell(row=row_input, column=col_idx).value
                    out_val = sheet.cell(row=row_output, column=col_idx).value
                    limits[(dept_name, model_base.lower(), current_section)] = (in_val, out_val)
                    
        return limits, dept_mapping
    except Exception as e:
        print(f"Error loading pricing database: {e}")
        return {}, {}

def lookup_token_limits(dept_name, model_name, limits, dept_mapping):
    dept_norm = dept_name.lower().strip()
    mapped_dept = dept_mapping.get(dept_norm)
    if not mapped_dept:
        for k in dept_mapping:
            if k in dept_norm:
                mapped_dept = dept_mapping[k]
                break
    if not mapped_dept:
        mapped_dept = "credit"
        
    model_lower = model_name.lower()
    model_base = None
    if "sonic" in model_lower:
        model_base = "sonic"
    elif "cheetah" in model_lower:
        model_base = "cheetah"
    elif "meticulous" in model_lower:
        model_base = "meticulous"
    else:
        model_base = "cheetah"
        
    session_type = "15m"
    if "12h" in model_lower or "12j" in model_lower or "12-hour" in model_lower:
        session_type = "12j"
    elif "24h" in model_lower or "24j" in model_lower or "24-hour" in model_lower:
        session_type = "24j"
        
    limit_vals = limits.get((mapped_dept, model_base, session_type))
    if not limit_vals:
        # standard defaults
        limit_vals = (209000, 1500)
        
    input_lim, output_lim = limit_vals
    
    if input_lim == "-" or input_lim is None:
        input_str = "-"
    else:
        try:
            input_str = f"{int(input_lim):,}".replace(",", ".")
        except:
            input_str = str(input_lim)
            
    if output_lim == "-" or output_lim is None:
        output_str = "-"
    else:
        try:
            output_str = f"{int(output_lim):,}".replace(",", ".")
        except:
            output_str = str(output_lim)
            
    return input_str, output_str

def get_time_threshold_seconds(model_name):
    model_name_lower = model_name.lower() if model_name else ""
    if "12h" in model_name_lower or "12j" in model_name_lower or "12-hour" in model_name_lower:
        return 12 * 3600
    elif "24h" in model_name_lower or "24j" in model_name_lower or "24-hour" in model_name_lower:
        return 24 * 3600
    else:
        return 15 * 60

def get_sheet_by_name(wb, name):
    name_clean = name.lower().replace(" ", "").strip()
    for sheet_name in wb.sheetnames:
        if sheet_name.lower().replace(" ", "").strip() == name_clean:
            return wb[sheet_name]
    raise KeyError(f"Worksheet '{name}' tidak ditemukan. Sheet yang tersedia di file ini: {wb.sheetnames}")

def parse_usage_excel(xlsx_path, pricing_xlsx_path):
    wb = openpyxl.load_workbook(xlsx_path, data_only=True)
    sheet_summary = get_sheet_by_name(wb, 'Summary')
    sheet_chat = get_sheet_by_name(wb, 'Chat Session')
    sheet_bill = get_sheet_by_name(wb, 'Billing Session')

    # 1. Read metadata
    client_name = sheet_summary.cell(row=12, column=3).value or 'PT Adira Dinamika Multi Finance Tbk'
    dept_name = sheet_summary.cell(row=13, column=3).value or 'Credit Data Analyst'
    period = sheet_summary.cell(row=14, column=3).value or '01 May 2026 – 31 May 2026'

    # 2. Read usage summary numbers
    chat_sessions = 0
    billing_sessions = 0
    input_tokens = 0
    output_tokens = 0
    
    for r in range(15, 30):
        val_b = sheet_summary.cell(row=r, column=2).value
        val_d = sheet_summary.cell(row=r, column=4).value
        if val_b == "CHAT SESSIONS":
            chat_sessions = sheet_summary.cell(row=r+1, column=2).value
        if val_d == "BILLING SESSIONS":
            billing_sessions = sheet_summary.cell(row=r+1, column=4).value
        if val_b == "TOTAL INPUT TOKENS":
            input_tokens = sheet_summary.cell(row=r+1, column=2).value
        if val_d == "TOTAL OUTPUT TOKENS":
            output_tokens = sheet_summary.cell(row=r+1, column=4).value

    # 3. Read models used dynamically
    models_used = []
    r_model = 18
    while True:
        no_val = sheet_summary.cell(row=r_model, column=2).value
        if not isinstance(no_val, int):
            break
        model_name = sheet_summary.cell(row=r_model, column=4).value
        if model_name and model_name not in models_used:
            models_used.append(model_name)
        r_model += 1

    # Load pricing limits db
    limits_db, dept_mapping = load_token_limits_db(pricing_xlsx_path)

    specs_list = []
    for m_name in models_used:
        in_lim, out_lim = lookup_token_limits(dept_name, m_name, limits_db, dept_mapping)
        window = "15 Menit"
        m_lower = m_name.lower()
        if "12h" in m_lower or "12j" in m_lower or "12-hour" in m_lower:
            window = "12 Jam"
        elif "24h" in m_lower or "24j" in m_lower or "24-hour" in m_lower:
            window = "24 Jam"
            
        specs_list.append({
            "model": m_name,
            "window": window,
            "input": in_lim,
            "output": out_lim
        })
    
    if not specs_list:
        specs_list = [{
            "model": "Cakra Cheetah - 15m",
            "window": "15 Menit",
            "input": "209.000",
            "output": "1.500"
        }]

    # 4. Analyze Over-billing
    chat_rows = list(sheet_chat.iter_rows(values_only=True))[7:]
    bill_rows = list(sheet_bill.iter_rows(values_only=True))[7:]

    billings_by_chat = {}
    for bs in bill_rows:
        if bs[1] is not None:
            billings_by_chat.setdefault(bs[3], []).append({
                'no': bs[1],
                'start_time': bs[2],
                'chat_uuid': bs[3],
                'bill_uuid': bs[4],
                'agent': bs[5],
                'model': bs[6],
                'logs': bs[7],
                'input_tokens': bs[8],
                'output_tokens': bs[9],
                'is_initial': bs[10]
            })

    over_billing_list = []
    count_time = 0
    count_output = 0
    count_both = 0
    max_bills = 1

    for r in chat_rows:
        if r[1] is not None and isinstance(r[1], int) and r[6] > 1:
            uuid = r[3]
            bills = billings_by_chat.get(uuid, [])
            
            if r[6] > max_bills:
                max_bills = r[6]

            has_time_limit = False
            has_output_limit = False
            
            for idx in range(1, len(bills)):
                prev_w = bills[idx-1]
                curr_w = bills[idx]
                
                # safety check if start_time is datetime
                if isinstance(curr_w['start_time'], datetime) and isinstance(prev_w['start_time'], datetime):
                    dt = (curr_w['start_time'] - prev_w['start_time']).total_seconds()
                else:
                    dt = 0
                
                threshold = get_time_threshold_seconds(prev_w['model'])
                if dt >= threshold:
                    has_time_limit = True
                else:
                    has_output_limit = True
            
            if has_time_limit and has_output_limit:
                count_both += 1
            elif has_time_limit:
                count_time += 1
            else:
                count_output += 1

            over_billing_list.append({
                'no': r[1],
                'uuid': uuid,
                'input_tokens': r[7],
                'output_tokens': r[8],
                'bills_count': r[6]
            })

    over_billing_count = len(over_billing_list)
    over_billing_pct = (over_billing_count / chat_sessions * 100) if chat_sessions else 0
    avg_billing_per_chat = billing_sessions / chat_sessions if chat_sessions else 0

    return {
        'client_name': client_name,
        'dept_name': dept_name,
        'period': period,
        'chat_sessions': chat_sessions,
        'billing_sessions': billing_sessions,
        'input_tokens': input_tokens,
        'output_tokens': output_tokens,
        'models_used': models_used,
        'specs_list': specs_list,
        'over_billing_count': over_billing_count,
        'over_billing_pct': over_billing_pct,
        'avg_billing_per_chat': avg_billing_per_chat,
        'count_time': count_time,
        'count_output': count_output,
        'count_both': count_both,
        'over_billing_list': over_billing_list,
        'xlsx_filename': os.path.basename(xlsx_path)
    }

def clean_json_response(text):
    text = text.strip()
    if text.startswith("```json"):
        text = text[7:]
    elif text.startswith("```"):
        text = text[3:]
    if text.endswith("```"):
        text = text[:-3]
    return text.strip()

def get_offline_fallback(data):
    sorted_over = sorted(data['over_billing_list'], key=lambda x: x['bills_count'], reverse=True)
    longest_sessions_parts = []
    for i in range(min(3, len(sorted_over))):
        longest_sessions_parts.append(f"UUID {sorted_over[i]['uuid'][:8]}... ({sorted_over[i]['bills_count']} Billing Sessions)")
    longest_sessions_str = ", ".join(longest_sessions_parts) if longest_sessions_parts else "-"

    return {
        "introduction": f"Merujuk pada implementasi layanan AI Chatbot pada aplikasi SYGMA, bersama ini kami sampaikan laporan ringkasan penggunaan layanan untuk divisi {data['dept_name']} pada periode {data['period']}. Laporan ini merupakan ikhtisar volume penggunaan, aktivitas interaksi chatbot, serta analisis detail billing session.",
        "insights": [
            f"Rata-rata Billing Session per Chat Session di divisi {data['dept_name']} adalah {data['avg_billing_per_chat']:.2f} sesi. Hal ini menandakan pola penggunaan interaktif dengan durasi percakapan melampaui batas awal sesi dasar.",
            f"Sesi dengan over-billing terbanyak terjadi pada: {longest_sessions_str}. Hal ini menandakan adanya sesi konsultasi intensif berdurasi sangat panjang atau melampaui batas limit token output berulang kali."
        ]
    }

def compile_docx(output_path, data, gemini_content, logo_path, enable_header=True):
    doc = Document()
    
    # Defaults
    style_normal = doc.styles['Normal']
    style_normal.font.name = 'Arial'
    style_normal.font.size = Pt(10)
    try:
        style_bullet = doc.styles['List Bullet']
        style_bullet.font.name = 'Arial'
        style_bullet.font.size = Pt(10)
    except:
        pass
        
    # Page setup
    section = doc.sections[0]
    section.page_width = Inches(8.27)
    section.page_height = Inches(11.69)
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)
    # 1. Logo Header
    if enable_header:
        p_logo = doc.add_paragraph()
        set_para_spacing(p_logo, before_pt=0, after_pt=10)
        if logo_path and os.path.exists(logo_path):
            p_logo.add_run().add_picture(logo_path, width=Inches(2.77))
        else:
            # fallback text logo if image is missing
            r = p_logo.add_run("ADIRA FINANCE")
            format_run(r, font_name="Arial", size_pt=18, bold=True, color_rgb=RGBColor(241, 196, 15))

        # Divider line
        p_div = doc.add_paragraph()
        set_para_spacing(p_div, before_pt=0, after_pt=10)
        p_div_border = parse_xml(f'<w:pBdr {nsdecls("w")}><w:bottom w:val="single" w:sz="6" w:space="1" w:color="1F4E79"/></w:pBdr>')
        p_div._p.get_or_add_pPr().append(p_div_border)

    # 2. Document Title
    p_title = doc.add_paragraph()
    set_para_spacing(p_title, before_pt=4, after_pt=12)
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_title = p_title.add_run("LAPORAN RINGKASAN PENGGUNAAN LAYANAN AI CHATBOT")
    format_run(run_title, font_name="Arial", size_pt=14, color_rgb=RGBColor(31, 78, 121), bold=True)

    # 3. Metadata Table (4 columns, borderless)
    table_meta = doc.add_table(rows=3, cols=4)
    table_meta.alignment = WD_TABLE_ALIGNMENT.CENTER
    table_meta.autofit = False
    set_table_borders(table_meta, color="FFFFFF", val="none")
    set_table_margins(table_meta, top=40, bottom=40, left=60, right=60)

    col_widths_meta = [Inches(1.1), Inches(2.28), Inches(0.8), Inches(2.59)]

    # Row 0
    p = table_meta.cell(0, 0).paragraphs[0]
    set_para_spacing(p, after_pt=2)
    r = p.add_run("Bulan Laporan")
    format_run(r, font_name="Arial", size_pt=9.5, bold=True)
    
    p = table_meta.cell(0, 1).paragraphs[0]
    set_para_spacing(p, after_pt=2)
    month_name = data['period'].split()[-1] if len(data['period'].split()) > 2 else "Laporan"
    r = p.add_run(f":  {month_name}")
    format_run(r, font_name="Arial", size_pt=9.5)
    
    p = table_meta.cell(0, 2).paragraphs[0]
    set_para_spacing(p, after_pt=2)
    r = p.add_run("Tanggal")
    format_run(r, font_name="Arial", size_pt=9.5, bold=True)
    
    p = table_meta.cell(0, 3).paragraphs[0]
    set_para_spacing(p, after_pt=2)
    today_str = datetime.now().strftime("%d %B %Y")
    # Translate month name to Indonesian simple map
    months_id = {
        "January": "Januari", "February": "Februari", "March": "Maret", "April": "April",
        "May": "Mei", "June": "Juni", "July": "Juli", "August": "Agustus",
        "September": "September", "October": "Oktoba", "November": "November", "December": "Desember"
    }
    for eng, ind in months_id.items():
        today_str = today_str.replace(eng, ind)
    r = p.add_run(f":  {today_str}")
    format_run(r, font_name="Arial", size_pt=9.5)

    # Row 1
    p = table_meta.cell(1, 0).paragraphs[0]
    set_para_spacing(p, after_pt=2)
    r = p.add_run("Perihal")
    format_run(r, font_name="Arial", size_pt=9.5, bold=True)
    
    p = table_meta.cell(1, 1).paragraphs[0]
    set_para_spacing(p, after_pt=2)
    r = p.add_run(f":  Laporan Ringkasan Penggunaan AI Chatbot {data['dept_name']}")
    format_run(r, font_name="Arial", size_pt=9.5)
    
    p = table_meta.cell(1, 2).paragraphs[0]
    set_para_spacing(p, after_pt=2)
    r = p.add_run("Kepada")
    format_run(r, font_name="Arial", size_pt=9.5, bold=True)
    
    p = table_meta.cell(1, 3).paragraphs[0]
    set_para_spacing(p, after_pt=2)
    r = p.add_run(f":  Yth. Tim {data['dept_name']}")
    format_run(r, font_name="Arial", size_pt=9.5)

    # Row 2
    p = table_meta.cell(2, 0).paragraphs[0]
    set_para_spacing(p, after_pt=2)
    r = p.add_run("Periode")
    format_run(r, font_name="Arial", size_pt=9.5, bold=True)
    
    p = table_meta.cell(2, 1).paragraphs[0]
    set_para_spacing(p, after_pt=2)
    r = p.add_run(f":  {data['period']}")
    format_run(r, font_name="Arial", size_pt=9.5)
    
    p = table_meta.cell(2, 2).paragraphs[0]
    set_para_spacing(p, after_pt=2)
    # empty cell
    
    p = table_meta.cell(2, 3).paragraphs[0]
    set_para_spacing(p, after_pt=2)
    r = p.add_run(f"   {data['client_name']}")
    format_run(r, font_name="Arial", size_pt=9.5)

    for row in table_meta.rows:
        for idx, w in enumerate(col_widths_meta):
            row.cells[idx].width = w

    # Spacer
    p_space = doc.add_paragraph()
    set_para_spacing(p_space, before_pt=0, after_pt=6)

    # 4. Salutation & Introduction (Gemini generated)
    p_salut = doc.add_paragraph()
    set_para_spacing(p_salut, before_pt=0, after_pt=6)
    r_salut = p_salut.add_run("Dengan hormat,")
    format_run(r_salut, font_name="Arial", size_pt=10)
    
    p_intro = doc.add_paragraph()
    p_intro.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    set_para_spacing(p_intro, before_pt=0, after_pt=6)
    
    intro_parts = gemini_content['introduction'].split('**')
    for idx, part in enumerate(intro_parts):
        if not part:
            continue
        r = p_intro.add_run(part)
        format_run(r, font_name="Arial", size_pt=10)
        if idx % 2 == 1:
            r.bold = True

    # 5. Section 1: Ringkasan Volume Penggunaan
    p_sec1 = doc.add_paragraph()
    set_para_spacing(p_sec1, before_pt=6, after_pt=4)
    r = p_sec1.add_run("1. Ringkasan Volume Penggunaan (Usage Overview)")
    format_run(r, font_name="Arial", size_pt=11, color_rgb=RGBColor(31, 78, 121), bold=True)

    table_usage = doc.add_table(rows=5, cols=3)
    table_usage.alignment = WD_TABLE_ALIGNMENT.CENTER
    table_usage.autofit = False
    set_table_borders(table_usage, color="D3D3D3")
    set_table_margins(table_usage, top=80, bottom=80, left=120, right=120)

    col_widths_usage = [Inches(2.0), Inches(1.3), Inches(3.47)]
    usage_rows = [
        ("Chat Sessions", format_id(data['chat_sessions']), "Sesi percakapan unik yang diinisiasi oleh pengguna."),
        ("Billing Sessions", format_id(data['billing_sessions']), "Sesi billing percakapan."),
        ("Total Input Tokens", format_id(data['input_tokens']), "Jumlah token input yang dikirim ke model AI."),
        ("Total Output Tokens", format_id(data['output_tokens']), "Jumlah token output yang dihasilkan oleh model AI.")
    ]

    headers_usage = ["Parameter Penggunaan", "Jumlah / Volume", "Keterangan"]
    hdr_row = table_usage.rows[0]
    for idx, text in enumerate(headers_usage):
        cell = hdr_row.cells[idx]
        cell.width = col_widths_usage[idx]
        set_cell_shading(cell, "1F4E79")
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_para_spacing(p, before_pt=0, after_pt=0)
        r = p.add_run(text)
        format_run(r, font_name="Arial", size_pt=9, color_rgb=RGBColor(255, 255, 255), bold=True)

    for r_idx, (param, val, desc) in enumerate(usage_rows, start=1):
        row = table_usage.rows[r_idx]
        cell_p = row.cells[0]
        cell_p.width = col_widths_usage[0]
        p_p = cell_p.paragraphs[0]
        set_para_spacing(p_p, before_pt=0, after_pt=0)
        r = p_p.add_run(param)
        format_run(r, font_name="Arial", size_pt=9, bold=True)
        
        cell_v = row.cells[1]
        cell_v.width = col_widths_usage[1]
        p_v = cell_v.paragraphs[0]
        p_v.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        set_para_spacing(p_v, before_pt=0, after_pt=0)
        r = p_v.add_run(val)
        format_run(r, font_name="Arial", size_pt=9)
        
        cell_d = row.cells[2]
        cell_d.width = col_widths_usage[2]
        p_d = cell_d.paragraphs[0]
        set_para_spacing(p_d, before_pt=0, after_pt=0)
        r = p_d.add_run(desc)
        format_run(r, font_name="Arial", size_pt=9)
        
        if r_idx % 2 == 0:
            set_cell_shading(cell_p, "F2F5F8")
            set_cell_shading(cell_v, "F2F5F8")
            set_cell_shading(cell_d, "F2F5F8")

    for row in table_usage.rows:
        for idx, w in enumerate(col_widths_usage):
            row.cells[idx].width = w

    # 6. Section 2: Model Specs
    p_sec2 = doc.add_paragraph()
    set_para_spacing(p_sec2, before_pt=10, after_pt=4)
    r = p_sec2.add_run("2. Spesifikasi Model LLM (Agents & AI Model Plans)")
    format_run(r, font_name="Arial", size_pt=11, color_rgb=RGBColor(31, 78, 121), bold=True)

    p_llm_intro = doc.add_paragraph()
    p_llm_intro.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    set_para_spacing(p_llm_intro, before_pt=0, after_pt=4)
    r = p_llm_intro.add_run(
        f"Layanan chatbot pada divisi {data['dept_name']} menggunakan agen AI dengan rincian limit token sebagai berikut:"
    )
    format_run(r, font_name="Arial", size_pt=10)

    table_llm = doc.add_table(rows=len(data['specs_list']) + 1, cols=4)
    table_llm.alignment = WD_TABLE_ALIGNMENT.CENTER
    table_llm.autofit = False
    set_table_borders(table_llm, color="D3D3D3")
    set_table_margins(table_llm, top=80, bottom=80, left=120, right=120)

    col_widths_llm = [Inches(2.2), Inches(1.4), Inches(1.5), Inches(1.57)]
    headers_llm = ["Model AI (Plan)", "Durasi Sesi", "Limit Input", "Limit Output"]

    hdr_row = table_llm.rows[0]
    for idx, text in enumerate(headers_llm):
        cell = hdr_row.cells[idx]
        cell.width = col_widths_llm[idx]
        set_cell_shading(cell, "1F4E79")
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_para_spacing(p, before_pt=0, after_pt=0)
        r = p.add_run(text)
        format_run(r, font_name="Arial", size_pt=9, color_rgb=RGBColor(255, 255, 255), bold=True)

    for r_idx, s_item in enumerate(data['specs_list'], start=1):
        row = table_llm.rows[r_idx]
        row_data = [s_item["model"], s_item["window"], s_item["input"], s_item["output"]]
        for c_idx, val in enumerate(row_data):
            cell = row.cells[c_idx]
            cell.width = col_widths_llm[c_idx]
            p = cell.paragraphs[0]
            set_para_spacing(p, before_pt=0, after_pt=0)
            
            if c_idx == 0:
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                r = p.add_run(val)
                format_run(r, font_name="Arial", size_pt=9, bold=True)
            elif c_idx == 1:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                r = p.add_run(val)
                format_run(r, font_name="Arial", size_pt=9)
            else:
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                r = p.add_run(val)
                format_run(r, font_name="Arial", size_pt=9)
                
            if r_idx % 2 == 0:
                set_cell_shading(cell, "F2F5F8")

    for row in table_llm.rows:
        for idx, w in enumerate(col_widths_llm):
            row.cells[idx].width = w

    # 7. Section 3: Summary of Over-billing
    p_sec3 = doc.add_paragraph()
    set_para_spacing(p_sec3, before_pt=10, after_pt=4)
    r = p_sec3.add_run("3. Ringkasan Analisis Sesi dengan Over-billing")
    format_run(r, font_name="Arial", size_pt=11, color_rgb=RGBColor(31, 78, 121), bold=True)

    p_over_intro = doc.add_paragraph()
    p_over_intro.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    set_para_spacing(p_over_intro, before_pt=0, after_pt=4)
    p_over_intro.add_run("Secara default, satu Chat Session dihitung sebagai satu Billing Session. Namun, sesi billing dapat bertambah (over-billing) apabila percakapan melebihi batas waktu sesi atau akumulasi token output menyentuh kapasitas model plan. Dari total ")
    r_chat_sess = p_over_intro.add_run(f"{format_id(data['chat_sessions'])}")
    format_run(r_chat_sess, font_name="Arial", size_pt=10)
    p_over_intro.add_run(" Chat Sessions di periode laporan, terdapat ")
    r_over_count = p_over_intro.add_run(f"{format_id(data['over_billing_count'])} sesi")
    format_run(r_over_count, font_name="Arial", size_pt=10, bold=True)
    p_over_intro.add_run(f" ({data['over_billing_pct']:.1f}%) yang mengalami over-billing. Berikut adalah ringkasan penyebab over-billing:")

    table_over_summary = doc.add_table(rows=5, cols=3)
    table_over_summary.alignment = WD_TABLE_ALIGNMENT.CENTER
    table_over_summary.autofit = False
    set_table_borders(table_over_summary, color="D3D3D3")
    set_table_margins(table_over_summary, top=80, bottom=80, left=120, right=120)

    col_widths_over_sum = [Inches(3.5), Inches(1.5), Inches(1.67)]
    headers_over_sum = ["Faktor Penyebab Over-billing", "Jumlah Sesi Chat", "Persentase (%)"]
    
    pct_time = (data['count_time'] / data['over_billing_count'] * 100) if data['over_billing_count'] else 0
    pct_output = (data['count_output'] / data['over_billing_count'] * 100) if data['over_billing_count'] else 0
    pct_both = (data['count_both'] / data['over_billing_count'] * 100) if data['over_billing_count'] else 0
    
    over_sum_rows = [
        ("Limit Token Output (Volume output menyentuh kapasitas)", data['count_output'], f"{pct_output:.1f}%"),
        ("Limit Waktu (Durasi percakapan / jeda melebihi threshold)", data['count_time'], f"{pct_time:.1f}%"),
        ("Kombinasi Limit Waktu & Token Output", data['count_both'], f"{pct_both:.1f}%"),
        ("TOTAL SESI OVER-BILLING", data['over_billing_count'], "100.0%")
    ]

    hdr_row = table_over_summary.rows[0]
    for idx, text in enumerate(headers_over_sum):
        cell = hdr_row.cells[idx]
        cell.width = col_widths_over_sum[idx]
        set_cell_shading(cell, "1F4E79")
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_para_spacing(p, before_pt=0, after_pt=0)
        r = p.add_run(text)
        format_run(r, font_name="Arial", size_pt=9, color_rgb=RGBColor(255, 255, 255), bold=True)

    for r_idx, (factor, count, pct) in enumerate(over_sum_rows, start=1):
        row = table_over_summary.rows[r_idx]
        
        cell_f = row.cells[0]
        cell_f.width = col_widths_over_sum[0]
        p = cell_f.paragraphs[0]
        set_para_spacing(p, before_pt=0, after_pt=0)
        r = p.add_run(factor)
        is_bold_row = (r_idx == 4)
        format_run(r, font_name="Arial", size_pt=9, bold=is_bold_row)
        
        cell_c = row.cells[1]
        cell_c.width = col_widths_over_sum[1]
        p = cell_c.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        set_para_spacing(p, before_pt=0, after_pt=0)
        r = p.add_run(format_id(count))
        format_run(r, font_name="Arial", size_pt=9, bold=is_bold_row)

        cell_p = row.cells[2]
        cell_p.width = col_widths_over_sum[2]
        p = cell_p.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        set_para_spacing(p, before_pt=0, after_pt=0)
        r = p.add_run(pct)
        format_run(r, font_name="Arial", size_pt=9, bold=is_bold_row)
        
        for col_idx in range(3):
            if r_idx == 4:
                set_cell_shading(row.cells[col_idx], "EAF2F8")
            elif r_idx % 2 == 0:
                set_cell_shading(row.cells[col_idx], "F2F5F8")

    for row in table_over_summary.rows:
        for idx, w in enumerate(col_widths_over_sum):
            row.cells[idx].width = w

    # 8. Section 4: General Insights (Gemini generated)
    p_sec4 = doc.add_paragraph()
    set_para_spacing(p_sec4, before_pt=10, after_pt=4)
    r = p_sec4.add_run("4. Ringkasan Analisis & Insight Penggunaan")
    format_run(r, font_name="Arial", size_pt=11, color_rgb=RGBColor(31, 78, 121), bold=True)

    for ins in gemini_content['insights']:
        # Clean the text: remove newlines, bullet signs, list numbers
        cleaned_ins = ins.replace('\n', ' ').strip()
        while '  ' in cleaned_ins:
            cleaned_ins = cleaned_ins.replace('  ', ' ')
        if cleaned_ins.startswith('* '):
            cleaned_ins = cleaned_ins[2:]
        elif cleaned_ins.startswith('- '):
            cleaned_ins = cleaned_ins[2:]
        elif cleaned_ins.startswith('1. ') or cleaned_ins.startswith('2. ') or cleaned_ins.startswith('3. ') or cleaned_ins.startswith('4. '):
            cleaned_ins = cleaned_ins[3:]
            
        p_ins = doc.add_paragraph(style='List Bullet')
        p_ins.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        set_para_spacing(p_ins, before_pt=0, after_pt=3)
        
        # Split by bold markers
        parts = cleaned_ins.split('**')
        for idx, part in enumerate(parts):
            if not part:
                continue
            r = p_ins.add_run(part)
            format_run(r, font_name="Arial", size_pt=10)
            if idx % 2 == 1:
                r.bold = True

    # 9. Section 5: Keterangan Lampiran & Penutup
    p_sec5 = doc.add_paragraph()
    set_para_spacing(p_sec5, before_pt=8, after_pt=4)
    r = p_sec5.add_run("5. Keterangan Lampiran")
    format_run(r, font_name="Arial", size_pt=11, color_rgb=RGBColor(31, 78, 121), bold=True)

    p_att = doc.add_paragraph()
    p_att.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_para_spacing(p_att, before_pt=0, after_pt=6)
    r_att = p_att.add_run(
        f"Rincian lengkap dari seluruh interaksi chatbot terlampir secara terpisah pada spreadsheet {data['xlsx_filename']}:\n"
        "1. Lampiran I (Chat Session): Histori chat lengkap per sesi percakapan unik.\n"
        "2. Lampiran II (Billing Session): Log detail per billing window interval untuk audit kepatuhan biaya."
    )
    format_run(r_att, font_name="Arial", size_pt=8.5, color_rgb=RGBColor(89, 89, 89), italic=True)

    p_close = doc.add_paragraph()
    p_close.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    set_para_spacing(p_close, before_pt=4, after_pt=12)
    r_close = p_close.add_run("Demikian laporan ringkasan penggunaan ini kami sampaikan. Atas perhatian dan kerjasamanya, kami ucapkan terima kasih.")
    format_run(r_close, font_name="Arial", size_pt=10)

    # 10. Section 6: Approval Block
    p_sec6 = doc.add_paragraph()
    set_para_spacing(p_sec6, before_pt=12, after_pt=4)
    r = p_sec6.add_run("6. Lembar Persetujuan Laporan (Approval Block)")
    format_run(r, font_name="Arial", size_pt=11, color_rgb=RGBColor(31, 78, 121), bold=True)

    table_appr = doc.add_table(rows=3, cols=5)
    table_appr.alignment = WD_TABLE_ALIGNMENT.CENTER
    table_appr.autofit = False
    set_table_borders(table_appr, color="999999")
    set_table_margins(table_appr, top=80, bottom=80, left=120, right=120)

    col_widths_appr = [Inches(1.8), Inches(1.3), Inches(1.6), Inches(1.1), Inches(0.97)]
    headers_appr = ["Role", "Name", "Comments", "Signature", "Date"]

    hdr_row = table_appr.rows[0]
    for idx, text in enumerate(headers_appr):
        cell = hdr_row.cells[idx]
        cell.width = col_widths_appr[idx]
        set_cell_shading(cell, "1F4E79")
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_para_spacing(p, before_pt=0, after_pt=0)
        r = p.add_run(text)
        format_run(r, font_name="Arial", size_pt=9, color_rgb=RGBColor(255, 255, 255), bold=True)

    appr_roles = [
        ("TMO Representative\n(PT Adira Dinamika Multi Finance Tbk)", "", "", "", ""),
        (f"{data['dept_name']} Representative\n(PT Adira Dinamika Multi Finance Tbk)", "", "", "", "")
    ]

    for r_idx, row_data in enumerate(appr_roles, start=1):
        row = table_appr.rows[r_idx]
        row.height = Inches(0.9)
        for c_idx, val in enumerate(row_data):
            cell = row.cells[c_idx]
            cell.width = col_widths_appr[c_idx]
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = cell.paragraphs[0]
            set_para_spacing(p, before_pt=0, after_pt=0)
            
            if c_idx == 0:
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                r = p.add_run(val)
                format_run(r, font_name="Arial", size_pt=9, bold=True)
            else:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                r = p.add_run("")
                format_run(r, font_name="Arial", size_pt=9)

    for row in table_appr.rows:
        for idx, w in enumerate(col_widths_appr):
            row.cells[idx].width = w

    doc.save(output_path)
